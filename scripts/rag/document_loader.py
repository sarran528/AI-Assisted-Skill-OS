from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from PyPDF2 import PdfReader


logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".html"}


@dataclass(slots=True)
class SourceDocument:
    skill_id: str
    phase: str | None
    technique_id: str | None
    doc_type: str
    source_path: str
    content: str


def _parse_filename(stem: str) -> tuple[str | None, str | None, str] | None:
    parts = stem.split("__")
    if len(parts) == 2:
        phase_slug, doc_type = parts
        technique_id = None
    elif len(parts) == 3:
        phase_slug, technique_id, doc_type = parts
        technique_id = technique_id or None
    else:
        logger.warning("Skipping document with invalid filename format: %s", stem)
        return None

    phase = None if phase_slug == "cross_phase" else phase_slug
    return phase, technique_id, doc_type


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _read_docx(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    return "\n".join(paragraphs)


def _read_html(path: Path) -> str:
    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "lxml")
    return soup.get_text(separator=" ")


def _read_file(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".txt":
        return _read_txt(path)
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".html":
        return _read_html(path)
    raise ValueError(f"Unsupported extension: {ext}")


def load_documents(skill_docs_dir: str) -> list[SourceDocument]:
    root = Path(skill_docs_dir)
    if not root.exists() or not root.is_dir():
        raise FileNotFoundError(f"Skill docs directory does not exist: {root}")

    documents: list[SourceDocument] = []

    for path in root.rglob("*"):
        if not path.is_file():
            continue

        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            logger.warning("Skipping unsupported document extension: %s", path)
            continue

        skill_id = path.parent.name
        parsed = _parse_filename(path.stem)
        if parsed is None:
            continue
        phase, technique_id, doc_type = parsed

        content = _read_file(path).strip()
        if not content:
            logger.warning("Skipping empty document: %s", path)
            continue

        documents.append(
            SourceDocument(
                skill_id=skill_id,
                phase=phase,
                technique_id=technique_id,
                doc_type=doc_type,
                source_path=str(path.relative_to(root).as_posix()),
                content=content,
            )
        )

    return documents
